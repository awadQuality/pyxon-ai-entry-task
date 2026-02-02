"""
Document Processing Service
Handles extraction of text from PDF, DOCX, and TXT files with full Arabic support
"""

import os
import re
from typing import Dict, Any, Tuple
from pathlib import Path
import chardet
from loguru import logger

# PDF Processing
import PyPDF2
import pdfplumber

# DOCX Processing
from docx import Document as DocxDocument


class DocumentProcessor:
    """Process documents and extract text with metadata"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document file and extract text with metadata
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        logger.info(f"Processing file: {file_path.name}")
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text, page_count = self.extract_pdf(str(file_path))
        elif file_extension in ['.docx', '.doc']:
            text, page_count = self.extract_docx(str(file_path))
        elif file_extension == '.txt':
            text, page_count = self.extract_txt(str(file_path))
        else:
            raise ValueError(f"Unsupported format: {file_extension}")
        
        # Detect language
        language = self.detect_language(text)
        
        # Validate Arabic diacritics if Arabic
        has_diacritics = False
        if language == 'arabic':
            has_diacritics = self.validate_arabic_diacritics(text)
        
        # Get file size
        file_size = file_path.stat().st_size
        
        metadata = {
            'text': text,
            'filename': file_path.name,
            'file_type': file_extension.replace('.', ''),
            'file_size': file_size,
            'language': language,
            'page_count': page_count,
            'has_arabic_diacritics': has_diacritics,
            'character_count': len(text),
            'word_count': len(text.split())
        }
        
        logger.info(f"Extracted {len(text)} characters from {file_path.name}")
        
        return metadata
    
    def extract_pdf(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF file
        Uses PyPDF2 as primary method with pdfplumber as fallback
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, page_count)
        """
        text = ""
        page_count = 0
        
        try:
            # Try PyPDF2 first (faster)
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If PyPDF2 didn't extract much text, try pdfplumber
            if len(text.strip()) < 100:
                logger.info("PyPDF2 extraction insufficient, trying pdfplumber...")
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
        
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
        
        return text.strip(), page_count
    
    def extract_docx(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (extracted_text, paragraph_count)
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text
            
            paragraph_count = len(paragraphs)
            
            return text.strip(), paragraph_count
        
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def extract_txt(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from TXT file with automatic encoding detection
        Supports UTF-8 and CP1256 (Arabic Windows encoding)
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Tuple of (extracted_text, line_count)
        """
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding']
            
            logger.info(f"Detected encoding: {encoding}")
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            line_count = len(text.split('\n'))
            
            return text.strip(), line_count
        
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            # Fallback to UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                return text.strip(), len(text.split('\n'))
            except:
                raise
    
    def detect_language(self, text: str) -> str:
        """
        Detect if text contains Arabic characters
        
        Args:
            text: Input text
            
        Returns:
            'arabic', 'english', or 'mixed'
        """
        # Arabic Unicode range: 0600-06FF
        arabic_pattern = re.compile(r'[\u0600-\u06FF]')
        english_pattern = re.compile(r'[a-zA-Z]')
        
        has_arabic = bool(arabic_pattern.search(text))
        has_english = bool(english_pattern.search(text))
        
        if has_arabic and has_english:
            return 'mixed'
        elif has_arabic:
            return 'arabic'
        elif has_english:
            return 'english'
        else:
            return 'unknown'
    
    def validate_arabic_diacritics(self, text: str) -> bool:
        """
        Check if Arabic text contains diacritics (harakat)
        
        Arabic diacritics Unicode ranges:
        - 064B-0652: Fatha, Damma, Kasra, Sukun, Shadda, etc.
        - 0670: Superscript Alef
        
        Args:
            text: Input text
            
        Returns:
            True if diacritics are present, False otherwise
        """
        diacritics_pattern = re.compile(r'[\u064B-\u0652\u0670]')
        return bool(diacritics_pattern.search(text))
